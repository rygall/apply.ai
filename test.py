from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Function to set paragraph spacing
def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1):
    p_pr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(space_before))
    spacing.set(qn('w:after'), str(space_after))
    spacing.set(qn('w:line'), str(line_spacing * 240))  # 240 is the default line height unit
    p_pr.append(spacing)


applicant_name = 'Ryan Gallagher'
applicant_town_and_state = 'Media, PA'
applicant_phone_number = '(610) 731-3822'
applicant_personal_email_address = 'ryan.gallagher900@gmail.com'
applicant_work_email_address = ''  # No work email address provided

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading(applicant_name, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(title, space_before=0, space_after=240)  # Add some space after the title

# Add a section for contact info
contact_info = doc.add_paragraph()
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(contact_info, space_before=0, space_after=240)  # Add some space after contact info
contact_info.add_run(applicant_town_and_state + ' | ').bold = True
contact_info.add_run(applicant_phone_number + ' | ')
contact_info.add_run(applicant_personal_email_address)

# Add a section for experience
experience_heading = doc.add_heading('Experience', level=1)
set_paragraph_spacing(experience_heading, space_before=240, space_after=120)

# Fill in applicant's experience
experience = [
    {
        'role': 'Mechanical Engineer',
        'company': 'Leidos, Philadelphia, PA',
        'period': 'May 2020 - Present',
        'place': '',
        'keypoints': [
            'Designed autonomous auxiliary systems for U.S. Navy and DARPA unmanned surface vessels.',
            'Created engine systems for U.S. Navy vessels, ensuring support for ship engines across all mission scenarios.',
            'Developed digital models for system performance evaluation and improvements.',
            'Provided programmatic and engineering support to the Naval Surface Warfare Center for the DDG(X) program.',
            'Analyzed data leading to optimized ship operations and performance enhancements.',
            'Developed software and machine learning models aiding ship design processes.'
        ]
    },
    {
        'role': 'Mechanical Engineering Intern',
        'company': 'Monroe Energy, LLC, Trainer, PA',
        'period': 'May - August 2016-2019',
        'place': '',
        'keypoints': []
    }
]

for exp in experience:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=60)
    p.add_run(f"{exp['role']} - {exp['company']}").bold = True
    p.add_run(f"\n{exp['period']} | {exp['place']}")
    for keypoint in exp['keypoints']:
        keypoint_paragraph = doc.add_paragraph(f"• {keypoint}", style='List Bullet')
        set_paragraph_spacing(keypoint_paragraph, space_before=0, space_after=0)

# Add a section for education
education_heading = doc.add_heading('Education', level=1)
set_paragraph_spacing(education_heading, space_before=240, space_after=120)

# Fill in applicant's education
education = [
    {
        'title': 'Master of Science, Computer Science',
        'university': 'Drexel University, Philadelphia, PA',
        'grad_date': '2021 - Present (Graduating Spring 2025)',
        'gpa': '3.93'
    },
    {
        'title': 'Bachelor of Science, Mechanical Engineering',
        'university': 'Thomas Jefferson University, Philadelphia, PA',
        'grad_date': '2018 - 2020',
        'gpa': '3.39'
    },
    {
        'title': 'Bachelor of Science, Physics',
        'university': 'West Chester University of Pennsylvania, West Chester, PA',
        'grad_date': '2015 - 2018',
        'gpa': '3.11',
        'minor': 'Mathematics'
    }
]

for edu in education:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=60)
    p.add_run(edu['title']).bold = True
    p.add_run(f"\n{edu['university']}")
    p.add_run(f"\nGraduation Date: {edu['grad_date']}")
    if 'gpa' in edu:
        p.add_run(f"\nGPA: {edu['gpa']}")
    if 'minor' in edu:
        p.add_run(f"\nMinor: {edu['minor']}")

# Add a section for skills
skills_heading = doc.add_heading('Skills', level=1)
set_paragraph_spacing(skills_heading, space_before=240, space_after=120)

skills = [
    {'area': 'Structural Engineering Support', 'description': 'Capable of providing technical support for project engineering packages and renovation projects.'},
    {'area': 'Technical Documentation', 'description': 'Experienced in creating comprehensive engineering documents, including scope of work, project drawings, and specification lists.'},
    {'area': 'Site Surveys', 'description': 'Skilled in performing site surveys to gather site-specific data and provide analysis.'},
    {'area': 'CAD Proficiency', 'description': 'Proficient in AutoCAD and other Computer-Aided Design tools such as SolidWorks and MATLAB.'},
    {'area': 'Engineering Software', 'description': 'Experience with NumPy, PyTorch, and Scikit-Learn for developing engineering solutions.'},
    {'area': 'Project Management', 'description': 'Familiar with coordinating and supporting large-scale infrastructure projects, including joint acceptance and contractor acceptance inspections.'},
    {'area': 'Engineering Analysis', 'description': 'Proficient in conducting and documenting engineering studies, cost-benefit analyses, and feasibility studies.'}
]

skills_paragraph = doc.add_paragraph()
set_paragraph_spacing(skills_paragraph, space_before=0, space_after=60)
skills_run = skills_paragraph.add_run()
for skill in skills:
    skills_run.add_text(f"{skill['area']}: {skill['description']}  |  ")

# Save the document
output_path = 'Resume.docx'
doc.save(output_path)

print(f"Document saved as {output_path}")