from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Function to set paragraph spacing
def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1):
    p_pr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(space_before))
    spacing.set(qn('w:after'), str(space_after))
    spacing.set(qn('w:line'), str(line_spacing * 240))  # 240 is the default line height unit
    p_pr.append(spacing)


applicant_name = 'FILL IN APPLICANT NAME'
applicant_town_and_state = 'FILL IN APPLICANT TOWN AND STATE'
applicant_phone_number = 'FILL IN APPLICANT NUMBER'
applicant_personal_email_address = 'FILL IN APPLICANT PERSONAL EMAIL ADDRESS'
applicant_work_email_address = 'FILL IN APPLICANT WORK EMAIL ADDRESS'


# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading(applicant_name, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(title, space_before=0, space_after=240)  # Add some space after the title

# Add a section for contact info
contact_info = doc.add_paragraph()
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(contact_info, space_before=0, space_after=240)  # Add some space after contact info
contact_info.add_run(applicant_town_and_state + ' | ').bold = True
contact_info.add_run(applicant_phone_number + ' | ')
contact_info.add_run(applicant_personal_email_address + ' | ')
contact_info.add_run(applicant_work_email_address)

# Add a section for experience
experience_heading = doc.add_heading('Experience', level=1)
set_paragraph_spacing(experience_heading, space_before=240, space_after=120)

# FILL IN APPLICANTS EXPERIENCE
experience = [
    {
    },
    {
    }
]

for exp in experience:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=60)
    p.add_run(f"{exp['role']} - {exp['company']}").bold = True
    p.add_run(f"\n{exp['period']} | {exp['place']}")
    for keypoint in exp['keypoints']:
        keypoint_paragraph = doc.add_paragraph(f"• {keypoint}", style='List Bullet')
        set_paragraph_spacing(keypoint_paragraph, space_before=0, space_after=0)

# Add a section for education
education_heading = doc.add_heading('Education', level=1)
set_paragraph_spacing(education_heading, space_before=240, space_after=120)

# FILL IN APPLICANTS EDUCATION
education = [
    {
    },
    {
    },
    {
    }
]

for edu in education:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=60)
    p.add_run(edu['title']).bold = True
    p.add_run(f"\n{edu['university']}")
    p.add_run(f"\nGraduation Date: {edu['grad_date']}")
    if 'gpa' in edu:
        p.add_run(f"\nGPA: {edu['gpa']}")
    if 'minor' in edu:
        p.add_run(f"\nMinor: {edu['minor']}")

# Add a section for skills
skills_heading = doc.add_heading('Skills', level=1)
set_paragraph_spacing(skills_heading, space_before=240, space_after=120)

skills = [
]

skills_paragraph = doc.add_paragraph()
set_paragraph_spacing(skills_paragraph, space_before=0, space_after=60)
skills_run = skills_paragraph.add_run()
for skill in skills:
    skills_run.add_text(f"{skill['area']}: {skill['description']}  |  ")

# Save the document
applicant_first = "FILL IN APPLICANT FIRST NAME HERE"
applicant_last = "FILL IN APPLICANT LAST NAME HERE"
document_path = "results\\" + applicant_first + "_" + applicant_last + "_Resume.docx"
doc.save(document_path)
print(f"Document saved as {document_path}")
