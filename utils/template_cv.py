from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_cover_letter():
    # Create a new Document
    doc = Document()

    # Add the header information
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_run = header.add_run("Name\nAddress\nCity, State, ZIP Code\nEmail Address\nPhone Number")
    header_run.font.size = Pt(12)

    # Add a blank line
    doc.add_paragraph()

    # Add the date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    date_run = date_paragraph.add_run("Date")
    date_run.font.size = Pt(12)

    # Add a blank line
    doc.add_paragraph()

    # Add the recipient's information
    recipient = doc.add_paragraph()
    recipient_run = recipient.add_run("Recipient Name\nRecipient Title\nCompany Name\nCompany Address\nCity, State, ZIP Code")
    recipient_run.font.size = Pt(12)

    # Add a blank line
    doc.add_paragraph()

    # Add the salutation
    salutation = doc.add_paragraph()
    salutation_run = salutation.add_run("FILL IN SALUTATION")
    salutation_run.font.size = Pt(12)

    # Add a blank line
    doc.add_paragraph()

    # Add the body of the cover letter
    body = doc.add_paragraph()
    body_run = body.add_run("FILL IN BODY OF COVER LETTER")
    body_run.font.size = Pt(12)

    # Add the closing
    closing = doc.add_paragraph()
    closing_run = closing.add_run("FILL IN CLOSING")
    closing_run.font.size = Pt(12)

    # Save the document
    applicant_first = "FILL IN APPLICANT FIRST NAME HERE"
    applicant_last = "FILL IN APPLICANT LAST NAME HERE"
    document_path = "results\\" + applicant_first + "_" + applicant_last + "_CV.docx"
    doc.save(document_path)
    print(f"Document saved as {document_path}")

create_cover_letter()